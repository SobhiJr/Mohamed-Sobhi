from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()
#Adding a profile pic
document.add_picture('wp2741597-american-psycho-wallpaper.jpg', width=Inches(2.0))

#Name, Phone number, Age details
speak('Hello user, What is your name? ')
name = input("What is your name? ")
speak('Hello '+ name + ' How are you today? ')
speak('What is your phone number? ')
phone_num = input('What is your phone number? ')
speak('What is your E-mail address? ')
email = input('What is your E-mail address? ')

document.add_paragraph(
    name + '\n' 
    + phone_num + '\n' 
    + email + '\n' 
)
#About section
document.add_heading('About me')
speak('Tell me about yourself? ')
about_me = input('Tell about yourself? ')
document.add_paragraph(about_me)

#Work Exp
document.add_heading('My experience')
p = document.add_paragraph()

speak('Can you enter the name of the companies you have worked at or your past experiences ? ')
company = input('Enter Company you have worked with ')
speak('The date of start? ')
started_date = input('From Date ')
speak('The date of the end? ')
ended_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(started_date + '-' + ended_date + '\n').italic = True

speak('Describe your experience at ' + company + '?  ' )
experience_details = input('Describe your experience at ' + company + '?  ' )
p.add_run(experience_details)

#More Exp
while True:
    speak('Is there is Other companies you have worked at? ')
    has_more_exps = input('Do you have more experiences? Yes or No?')
    if has_more_exps.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company you have worked with ')
        speak('The date of start? ')
        started_date = input('From Date ')
        speak('The date of the end? ')
        ended_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(started_date + '-' + ended_date + '\n').italic = True

        experience_details = input('Describe your experience at ' + company + '?  ' )
        p.add_run(experience_details)
    else:
        break

#Skills
document.add_heading('Skills')
speak('Enter your unique skills ')
skill = input('Enter your unique skills ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    speak('Do you have any other unique skills? ')
    has_more_skills = input('Do you have more skills? Yes or No? ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter your uniqe skills ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

speak('Thanks for using this project, your CV is ready now')
#Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated by Mohamed Sobhi"

document.save('CV.docx')
